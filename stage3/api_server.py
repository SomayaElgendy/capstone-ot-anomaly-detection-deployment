from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from typing import Optional
import json
import tempfile
from pathlib import Path
from fastapi.responses import Response, StreamingResponse
from io import BytesIO
from docx import Document

from run_deployment import run_pipeline, run_chat

app = FastAPI(title="Stage 3 IR Service")


class AlertRequest(BaseModel):
    predicted_attack: str
    classifier_confidence: float
    network_anomaly_score: float
    process_anomaly_score: float
    window_start_time: str
    window_end_time: str
    technique_id: Optional[str] = None


class ChatRequest(BaseModel):
    message: str
    context: dict


@app.get("/health")
def health():
    return {"status": "ok"}


def find_ir_file_by_run_id(run_id: str) -> Path:
    ir_dir = Path("outputs/stage3/IR")
    ir_path = ir_dir / f"{run_id}.md"

    if not ir_path.exists():
        raise HTTPException(status_code=404, detail=f"IR report not found for run_id: {run_id}")

    return ir_path


def markdown_to_docx_stream(markdown_text: str, run_id: str) -> BytesIO:
    document = Document()

    document.add_heading("Incident Response Report", level=1)
    document.add_paragraph(f"Run ID: {run_id}")

    for line in markdown_text.splitlines():
        stripped = line.strip()

        if not stripped:
            continue

        if stripped.startswith("# "):
            document.add_heading(stripped.replace("# ", "", 1), level=1)

        elif stripped.startswith("## "):
            document.add_heading(stripped.replace("## ", "", 1), level=2)

        elif stripped.startswith("**") and stripped.endswith("**"):
            document.add_paragraph(stripped.replace("**", ""))

        elif stripped.startswith("- "):
            document.add_paragraph(stripped.replace("- ", "", 1), style="List Bullet")

        elif stripped[:3].replace(".", "").isdigit() and ". " in stripped[:5]:
            document.add_paragraph(stripped, style="List Number")

        else:
            document.add_paragraph(stripped.replace("**", ""))

    stream = BytesIO()
    document.save(stream)
    stream.seek(0)
    return stream

@app.post("/generate-ir")
def generate_ir(alert: AlertRequest):
    try:
        payload = alert.model_dump()

        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir) / "temp_alert.json"

            with open(tmp_path, "w", encoding="utf-8") as f:
                json.dump([payload], f, indent=2)

            result = run_pipeline(alerts_path=str(tmp_path))

        if not result:
            raise HTTPException(status_code=500, detail="No result returned from pipeline")

        if not result.get("success", False):
            raise HTTPException(
                status_code=500,
                detail=result.get("error", "IR generation failed"),
            )

        report = result.get("report", "")

        if not report:
            raise HTTPException(status_code=500, detail="Empty report returned")

        return {
            "reply": report,
            "grade": result.get("grade"),
            "score": result.get("score"),
            "run_id": result.get("run_id"),
            "ir_path": result.get("ir_path"),
            "eval_path": result.get("eval_path"),
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/chat")
def chat_endpoint(req: ChatRequest):
    try:
        result = run_chat(message=req.message, context=req.context)

        if not result:
            raise HTTPException(status_code=500, detail="No result returned from chat pipeline")

        if not result.get("success", False):
            raise HTTPException(
                status_code=500,
                detail=result.get("error", "Chat generation failed"),
            )

        reply = result.get("reply", "")

        if not reply:
            raise HTTPException(status_code=500, detail="Empty reply returned")

        return {
            "reply": reply,
            "generation_time_s": result.get("generation_time_s"),
            "finish_reason": result.get("finish_reason"),
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
@app.get("/download-ir/{run_id}/md")
def download_ir_markdown(run_id: str):
    ir_path = find_ir_file_by_run_id(run_id)
    markdown_text = ir_path.read_text(encoding="utf-8")

    return Response(
        content=markdown_text,
        media_type="text/markdown",
        headers={
            "Content-Disposition": f'attachment; filename="{run_id}.md"'
        },
    )

@app.get("/download-ir/{run_id}/docx")
def download_ir_docx(run_id: str):
    ir_path = find_ir_file_by_run_id(run_id)
    markdown_text = ir_path.read_text(encoding="utf-8")

    docx_stream = markdown_to_docx_stream(markdown_text, run_id)

    return StreamingResponse(
        docx_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{run_id}.docx"'
        },
    )